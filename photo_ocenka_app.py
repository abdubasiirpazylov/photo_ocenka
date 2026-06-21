import streamlit as st
from docx import Document
from docx.shared import Mm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import io
from PIL import Image, ImageOps

st.set_page_config(page_title="Эксперт - Фотоотчет", layout="wide", initial_sidebar_state="collapsed")

st.title("📸 Быстрый Фотоотчет")
st.markdown("Загрузите фото с места осмотра. Приложение автоматически соберет таблицу для отчета.")

reg_num = st.text_input("Гос. номер автомобиля (для названия файла):", placeholder="Например: 04KG021AJF")

st.header("Фотографии")

CAPTION_OPTIONS = [
    "Вид спереди", "Вид сзади", "Вид слева", "Вид справа",
    "VIN-код", "Показания одометра", "Обзорный снимок салона",
    "вмятина", "вытяжение", "гофра", "деформация", "изгиб", 
    "повреждение ЛКП", "потертость", "разрушение", "разрыв", 
    "раскол", "складка", "царапина", "скол", "перекос",
    "без образования видимых складок", "с глубокой вытяжкой металла", 
    "с нарушением геометрии кромок", "с нарушением геометрии ребер жесткости", 
    "с незначительной вытяжкой металла", "с образованием незначительных складок", 
    "с образованием острых складок", "с образованием плавных складок", "с образованием трещин",
    "на площади менее 10%", "на площади от 10 до 20%", 
    "на площади от 20 до 30%", "на площади от 30 до 40%", "на площади более 40%"
]

if "photo_order" not in st.session_state:
    st.session_state.photo_order = []  
if "processed_ids" not in st.session_state:
    st.session_state.processed_ids = set() 
if "photo_data" not in st.session_state:
    st.session_state.photo_data = {} 

# --- УМНАЯ ОБРЕЗКА ФОТО 1:1 И ВЫРАВНИВАНИЕ ---
def process_and_crop_image(uploaded_file):
    image = Image.open(uploaded_file)
    image = ImageOps.exif_transpose(image) # Выравниваем горизонт
    
    if image.mode in ("RGBA", "P"):
        image = image.convert("RGB")
        
    # --- Делаем идеальный квадрат (1:1) по центру ---
    width, height = image.size
    min_dim = min(width, height)
    left = (width - min_dim) / 2
    top = (height - min_dim) / 2
    right = (width + min_dim) / 2
    bottom = (height + min_dim) / 2
    image = image.crop((left, top, right, bottom))
    # ------------------------------------------------
    
    img_byte_arr = io.BytesIO()
    image.save(img_byte_arr, format='JPEG', quality=85)
    img_byte_arr.seek(0)
    return img_byte_arr

# --- КОЛБЭКИ ---
def move_photo(file_id, current_index):
    new_val = st.session_state.get(f"pos_{file_id}")
    if new_val is None:
        return
    new_index = new_val - 1
    if new_index != current_index:
        item = st.session_state.photo_order.pop(current_index)
        st.session_state.photo_order.insert(new_index, item)
        for idx, fid in enumerate(st.session_state.photo_order):
            st.session_state[f"pos_{fid}"] = idx + 1

def delete_photo(file_id):
    if file_id in st.session_state.photo_order:
        st.session_state.photo_order.remove(file_id)
    if file_id in st.session_state.processed_ids:
        st.session_state.processed_ids.remove(file_id)
    if file_id in st.session_state.photo_data:
        del st.session_state.photo_data[file_id]
        
    for idx, fid in enumerate(st.session_state.photo_order):
        st.session_state[f"pos_{fid}"] = idx + 1

# --- ЗАГРУЗКА ---
uploaded_photos = st.file_uploader("Сделать снимок или выбрать из галереи", type=["jpg", "jpeg", "png"], accept_multiple_files=True)

if uploaded_photos:
    for uf in uploaded_photos:
        if uf.file_id not in st.session_state.processed_ids:
            # Вызываем нашу новую функцию с обрезкой 1:1
            fixed_bytes = process_and_crop_image(uf)
            
            st.session_state.processed_ids.add(uf.file_id)
            st.session_state.photo_data[uf.file_id] = fixed_bytes
            st.session_state.photo_order.append(uf.file_id)
            st.session_state[f"pos_{uf.file_id}"] = len(st.session_state.photo_order)

photo_data_list = []

if st.session_state.photo_order:
    st.write("Настройте подписи и порядок фотографий:")
    
    for i, file_id in enumerate(list(st.session_state.photo_order)):
        img_bytes = st.session_state.photo_data[file_id]
        
        with st.container():
            c_img, c_controls = st.columns([1, 2])
            
            with c_img:
                st.image(img_bytes, use_container_width=True)
                st.selectbox(
                    "📍 Позиция:",
                    options=list(range(1, len(st.session_state.photo_order) + 1)),
                    key=f"pos_{file_id}",
                    on_change=move_photo,
                    args=(file_id, i)
                )
                st.button("❌ Исключить", key=f"del_{file_id}", on_click=delete_photo, args=(file_id,))

            with c_controls:
                selected_tags = st.multiselect("Шаблонные фразы:", CAPTION_OPTIONS, key=f"tags_{file_id}")
                custom_text = st.text_input("Свой текст:", key=f"custom_{file_id}")
                
                final_caption_parts = []
                if selected_tags:
                    final_caption_parts.append(", ".join(selected_tags))
                if custom_text:
                    final_caption_parts.append(custom_text)
                    
                final_caption = ", ".join(final_caption_parts)
                
                if final_caption:
                    st.caption(f"📝 Подпись: **{final_caption}**")
                    
            st.divider() 
            
            photo_data_list.append({"bytes": img_bytes, "caption": final_caption})

# --- ГЕНЕРАЦИЯ ОТЧЕТА ---
if photo_data_list:
    if st.button("СГЕНЕРИРОВАТЬ ФОТООТЧЕТ", type="primary", use_container_width=True):
        try:
            doc = Document()
            
            # Заголовок удален!
            
            # Создаем таблицу
            table = doc.add_table(rows=0, cols=2)
            table.style = 'Table Grid'
            
            for i in range(0, len(photo_data_list), 2):
                cells = table.add_row().cells
                
                # --- ЛЕВАЯ ЯЧЕЙКА ---
                cells[0].vertical_alignment = WD_ALIGN_VERTICAL.CENTER # Выравниваем ячейку по вертикали
                p1 = cells[0].paragraphs[0]
                p1.alignment = WD_ALIGN_PARAGRAPH.CENTER # Выравниваем по центру
                p1.paragraph_format.space_after = Pt(0) # Убираем лишние отступы снизу
                
                img1_file = photo_data_list[i]["bytes"]
                img1_file.seek(0)
                
                # Вставляем фото (ширина 80мм, так как фото 1:1, высота тоже будет 80мм)
                run1_img = p1.add_run()
                run1_img.add_picture(img1_file, width=Mm(80))
                
                # Вставляем текст
                run1_text = p1.add_run("\n" + photo_data_list[i]["caption"])
                run1_text.font.name = 'Cambria' # Устанавливаем шрифт
                run1_text.font.size = Pt(8)     # Устанавливаем размер 8
                
                # --- ПРАВАЯ ЯЧЕЙКА ---
                if i + 1 < len(photo_data_list):
                    cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                    p2 = cells[1].paragraphs[0]
                    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p2.paragraph_format.space_after = Pt(0)
                    
                    img2_file = photo_data_list[i+1]["bytes"]
                    img2_file.seek(0)
                    
                    run2_img = p2.add_run()
                    run2_img.add_picture(img2_file, width=Mm(80))
                    
                    run2_text = p2.add_run("\n" + photo_data_list[i+1]["caption"])
                    run2_text.font.name = 'Cambria'
                    run2_text.font.size = Pt(8)
            
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            st.success("✅ Отчет успешно сгенерирован!")
            
            safe_reg_num = reg_num.strip() if reg_num.strip() else "Без_номера"
            file_name = f"Фотоотчет_{safe_reg_num}.docx"
            
            st.download_button(
                label=f"📥 СКАЧАТЬ {file_name}",
                data=buffer,
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
            
        except Exception as e:
            st.error(f"Произошла ошибка при генерации отчета: {e}")
